from fastapi.responses import Response
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from datetime import datetime, timedelta
from docx.shared import Inches, Pt, Cm
from Model.model import ReportRequest
from utils.upload_to_s3 import get_report_from_s3_function
from word_report_generator.utils import (
    add_header_footer_images,
    add_key_value,
    add_bullet_points,
    add_table_row,
    add_paragraph,
    create_table,
    add_bullet,
    add_underline_line
)

ASSETS_DIR = "assets/images"
header_img = f"{ASSETS_DIR}/full_header.png"
footer_img = f"{ASSETS_DIR}/footer.png"

# === Helper Functions ===


def get_section_number(increment=True):
    global SECTION_NO
    if increment:
        SECTION_NO += 1
    return SECTION_NO

def add_title_with_number(
    doc: Document,
    text: str,
    size: int = 24,
    bold: bool = True,
    center: bool = False,
    numbered: bool = True,
    level: int = 1,
):
    """Add a section title (Heading 1/2/3) that supports TOC."""
    if numbered:
        section_number = get_section_number(increment=True)
        text = f"{section_number}. {text}"

    style_name = f"Heading {level}"

    para = doc.add_paragraph(text, style=style_name)
    run = para.runs[0]
    run.font.size = Pt(size)
    run.bold = bold

    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
def parse_public_access_block(pab_data):
    """
    parses the 'public_access_block_configuration' field from S3 bucket scan results.
    """
    pab_summary = []

    if pab_data is None:
        return "All settings disabled (false)"

    if isinstance(pab_data, list):
        for cfg in pab_data:
            if isinstance(cfg, dict):
                for k, v in cfg.items():
                    pab_summary.append(f"{k}: {'true' if v else 'false'}")
            else:
                pab_summary.append(str(cfg))

    elif isinstance(pab_data, dict):
        for k, v in pab_data.items():
            pab_summary.append(f"{k}: {'true' if v else 'false'}")

    elif isinstance(pab_data, str):
        pab_summary.append(pab_data)

    else:
        pab_summary.append("All settings disabled (false)")

    return "; ".join(pab_summary) if pab_summary else "All settings disabled (false)"


def add_no_data_note(doc: Document, text: str, size: int = 12, bold: bool = False):
    """
    Adds a simple note to the document, used when no data/resources are found.
    """
    para = doc.add_paragraph()
    run = para.add_run(text)

    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold

    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

    return para


# === section functions ===


def add_section_executive_summary(doc, data):
    """
    Adds the Executive Summary section
    """

    add_title_with_number(doc, "Executive Summary", size=18)

    exec_summary_text = (
        "This AWS Security Assessment was conducted to evaluate the current security posture "
        "of the client’s AWS environment. The assessment is based on AWS security best practices and CloudThat Technologies’ internal checklists."
    )
    add_paragraph(doc, exec_summary_text)

    # Initialize counters
    high_count = 0
    medium_count = 0
    low_count = 0

    # Helper function to update severity counters
    def accumulate_from_check(check_info):
        nonlocal high_count, medium_count, low_count
        sev = str(check_info.get("severity_level", "")).lower()
        if sev == "high":
            high_count += 1
        elif sev == "medium":
            medium_count += 1
        elif sev == "low":
            low_count += 1

    # Parse global services scan results
    global_results = data.get("global_services_scan_results", {})
    for key, check_info in global_results.items():
        if isinstance(check_info, dict):
            accumulate_from_check(check_info)

    # Parse regional results
    regional_results = data.get("results", [])
    for region_block in regional_results:
        region_data = region_block.get("data", {})
        for check_key, check_info in region_data.items():
            if isinstance(check_info, dict):
                accumulate_from_check(check_info)

    total_findings = high_count + medium_count + low_count

    # Add summary bullets
    add_key_value(doc, "Summary of Key Observations", "")
    bullet_points = [
        f"Total Security Findings: {total_findings}",
        f"High-Risk Issues: {high_count}",
        f"Medium-Risk Issues: {medium_count}",
        f"Low-Risk Issues: {low_count}",
    ]
    add_bullet_points(doc, bullet_points)
    # add_underline_line(doc)


def add_section_assessment_scope(doc, data):

    add_title_with_number(doc, "Assessment Scope", size=18)

    # Create table for Assessment Scope
    scope_table = create_table(doc, ["Parameter", "Details"])

    services_list = sorted(
        set(
            service
            for meta in data.get("scanned_meta_data", [])
            for service in meta.get("data", {}).get("services_scanned", [])
        )
    )
    regions_list = data.get("regions", [])

    security_services_list = (
        sorted(
            set(
                service
                for region_data in data.get("security_services_scanned_data", [])
                for service in region_data.get("data", {}).keys()
            )
        )
        if data.get("security_services_scanned_data")
        else []
    )

    # Add rows
    add_table_row(scope_table, ["AWS Services Assessed", ", ".join(services_list)])
    add_table_row(scope_table, ["AWS Regions Covered", ", ".join(regions_list)])
    add_table_row(
        scope_table, ["Security Services Reviewed", ", ".join(security_services_list)]
    )
    add_table_row(
        scope_table,
        [
            "Compliance Standards Referenced",
            "AWS Security Best Practices",
        ],
    )

    # add_underline_line(doc)


def add_section_security_services_status(doc, data):
    """
    Adds the Security Services Status section with a table.
    - Columns: Service, Region, Status, Comments
    """

    add_title_with_number(doc, "Security Services Status", size=18)

    headers = ["Service", "Region", "Status", "Comments"]
    table = create_table(doc, headers)

    grouped = {}

    for region_block in data.get("security_services_scanned_data", []):
        try:
            region = region_block.get("region", "-")
            region_services = region_block.get("data", {}) or {}
        except Exception:
            continue

        for service_name, service_info in region_services.items():
            try:
                # Prepare storage list
                if service_name not in grouped:
                    grouped[service_name] = []

                # NULL entry -> default values
                if service_info is None:
                    grouped[service_name].append(
                        (
                            region,
                            "Not Enabled",
                            "Service data unavailable or service not enabled",
                        )
                    )
                    continue

                # Extract enable status
                enabled = str(service_info.get("is_enabled", "")).lower() == "yes"

                status = "Enabled" if enabled else "Not Enabled"
                comments = (
                    "N/A"
                    if enabled
                    else service_info.get(
                        "recommendation", "No recommendation available"
                    )
                )

                grouped[service_name].append((region, status, comments))

            except Exception as inner_err:
                print(f"Error processing service '{service_name}': {inner_err}")
                continue

    # add rows to table WITH MERGED FIRST COLUMN
    for service_name, rows in grouped.items():
        try:
            if not rows:
                continue

            start_idx = len(table.rows)

            # Insert all rows
            for i, (region, status, comments) in enumerate(rows):
                first_col = service_name if i == 0 else ""
                add_table_row(table, [first_col, region, status, comments])

            end_idx = len(table.rows) - 1

            # Merge first column if multiple rows
            if len(rows) > 1:
                try:
                    top_cell = table.rows[start_idx].cells[0]
                    bottom_cell = table.rows[end_idx].cells[0]
                    merged = top_cell.merge(bottom_cell)
                    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as merge_err:
                    print(
                        f"Warning: failed to merge cells for '{service_name}': {merge_err}"
                    )

        except Exception as service_err:
            print(f"Error finalizing service '{service_name}': {service_err}")
            continue


def add_section_iam(doc, data):
    """
    Adds the Identity and Access Management (IAM) section.
    """

    add_title_with_number(doc, "Identity and Access Management (IAM)", size=18)

    # Create table
    headers = ["Check", "Status", "Observation", "Recommendation"]
    table = create_table(doc, headers)

    global_checks = data.get("global_services_scan_results", {})

    for check_key, item in global_checks.items():

        if item.get("service", "").lower() != "iam":
            continue

        check_name = item.get("check_name", check_key)
        additional_info = item.get("additional_info", {})
        affected = additional_info.get("affected", 0)
        resources = item.get("resources_affected", [])
        recommendation_text = item.get(
            "recommendation", "Review findings and apply best practices"
        )

        #  Status
        status = "Found" if affected and affected > 0 else "Not Found"

        #  Observation
        if status == "Found":
            resource_names = []
            for r in resources:
                if isinstance(r, dict):
                    name = (
                        r.get("resource_name")
                        or r.get("resource_id")
                        or r.get("user_id")
                    )
                    if name:
                        resource_names.append(name)
                elif isinstance(r, str):
                    resource_names.append(r)

            example_part = ""
            if resource_names:
                examples = resource_names[:4]
                example_part = f"(e.g., {', '.join(examples)}" + (
                    ", and more)" if len(resource_names) > 4 else ")"
                )

            observation = f"{affected} {check_name} {example_part}"
        else:
            observation = "No issues detected"

        #  Recommendation
        recommendation = recommendation_text if status == "Found" else "N/A"

        # Add row
        add_table_row(table, [check_name, status, observation, recommendation])

    # add_underline_line(doc)


def add_section_s3_bucket_security(doc, data):
    """
    Adds the S3 Bucket Security section
    """

    #  Title
    add_title_with_number(doc, "S3 Bucket Security", size=18)

    #  Extract S3 Bucket Data
    s3_buckets = (
        data.get("global_services_scan_results", {})
        .get("public_s3_buckets", {})
        .get("resources_affected", [])
    )
    if not s3_buckets:
        add_no_data_note(doc, "No public S3 buckets found.")
        return

    #  Table Headers
    headers = ["Bucket Name", "Public Access", "Action Needed"]
    table = create_table(doc, headers)

    for bucket in s3_buckets:
        bucket_name = bucket.get("resource_name", "Unknown")

        # Determine if public access is enabled
        pab_configs = bucket.get("public_access_block_configuration", [])
        pab_text = "Yes" if pab_configs else "No"

        # Define action based on public access
        action_needed = (
            "Check & Block Public Access" if pab_text == "Yes" else "No Action Needed"
        )

        add_table_row(table, [bucket_name, pab_text, action_needed])
    # add_underline_line(doc)


def add_section_ec2_network_security(doc, data):
    """
    Adds the EC2 and Network Security section.
    Columns: Check, Region, Findings, Recommendations
    """

    add_title_with_number(doc, "EC2 and Network Security", size=18)

    headers = ["Check", "Region", "Findings", "Recommendations"]
    table = create_table(doc, headers)

    # Mapping internal keys to readable titles
    check_mapping = {
        "open_security_groups": "Insecure Security Groups",
        "unused_security_groups": "Unused Security Groups",
        "unencrypted_ebs_volumes": "EBS Volume Encryption",
    }

    # Collect rows grouped by check
    grouped = {display: [] for display in check_mapping.values()}

    for region_block in data.get("results", []):
        region = region_block.get("region", "-")
        region_data = region_block.get("data", {}) or {}

        for key, display_name in check_mapping.items():
            try:
                check_info = region_data.get(key)
            except Exception:
                continue

            if not check_info:
                continue

            # Findings
            resources = check_info.get("resources_affected", []) or []
            total_affected = len(resources)
            findings_text = (
                f"{total_affected} {display_name.lower()}" if total_affected else "None"
            )

            # Recommendation
            recommendation = (
                check_info.get("recommendation", "Review and apply best practices")
                if total_affected > 0
                else "N/A"
            )

            # Store grouped row
            grouped[display_name].append((region, findings_text, recommendation))

    # Build table with merged first column
    for check_name, rows in grouped.items():
        try:
            if not rows:
                continue

            start_idx = len(table.rows)

            for i, (region, findings, rec) in enumerate(rows):
                first_col_value = check_name if i == 0 else ""
                add_table_row(table, [first_col_value, region, findings, rec])

            end_idx = len(table.rows) - 1

            if len(rows) > 1:
                try:
                    top_cell = table.rows[start_idx].cells[0]
                    bottom_cell = table.rows[end_idx].cells[0]
                    merged = top_cell.merge(bottom_cell)
                    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as merge_err:
                    print(
                        f"Warning: failed to merge cells for '{check_name}': {merge_err}"
                    )

        except Exception as chk_err:
            print(f"Error processing check '{check_name}': {chk_err}")
            continue


def add_report_header_section(doc, data):
    #  Add Title
    add_title_with_number(doc, "AWS Security Assessment Report", size=24)

    #  Key Info Section
    today = datetime.today().strftime("%d/%m/%y")
    last_7_days = (datetime.today() - timedelta(days=7)).strftime("%d/%m/%y")

    add_key_value(doc, "Prepared by", "CloudThat Technologies")
    add_key_value(doc, "Client Name", data.get("account_name", "N/A"))
    add_key_value(doc, "AWS Account ID(s)", data.get("account_id", "N/A"))
    add_key_value(doc, "Assessment Period", f"{last_7_days} to {today}")
    add_key_value(doc, "Prepared Date", today)
    add_key_value(doc, "Prepared By", "")
    # add_underline_line(doc)


def add_section_guardduty_findings(doc, data):
    """
    7. Threat Detection - GuardDuty Findings
    """

    add_title_with_number(doc, "Threat Detection - GuardDuty Findings", size=18)

    headers = ["Severity", "Region", "Findings Count"]
    table = create_table(doc, headers)

    # Severity mapping
    severity_ranges = {
        "Low": {"gte": 0.1, "lt": 4.0},
        "Medium": {"gte": 4.0, "lt": 7.0},
        "High": {"gte": 7.0, "lt": 9.0},
        "Critical": {"gte": 9.0, "lte": 10.0},
    }

    # Prepare grouped storage exactly like EC2 code
    severities_order = ["Critical", "High", "Medium", "Low"]
    grouped = {sev: [] for sev in severities_order}

    regional_results = data.get("results", [])

    # Iterate every region and count findings
    for region_block in regional_results:
        region = region_block.get("region", "-")
        region_data = region_block.get("data", {}) or {}

        guardduty_data = region_data.get("guardduty_findings", {})
        findings = guardduty_data.get("resources_affected", [])

        # Count per severity
        severity_counts = {"Low": 0, "Medium": 0, "High": 0, "Critical": 0}

        for f in findings:
            sev = float(f.get("severity", 0))

            for sev_label, r in severity_ranges.items():
                gte = r["gte"]
                lt = r.get("lt")
                lte = r.get("lte")

                if lt and gte <= sev < lt:
                    severity_counts[sev_label] += 1
                elif lte and gte <= sev <= lte:
                    severity_counts[sev_label] += 1

        # Append to grouping
        for sev_label in severities_order:
            grouped[sev_label].append((region, str(severity_counts[sev_label])))

    # Build table with merged severity column
    for severity_name, rows in grouped.items():
        try:
            if not rows:
                continue

            start_idx = len(table.rows)

            # Add row for every region
            for i, (region, count) in enumerate(rows):
                first_col_value = severity_name if i == 0 else ""
                add_table_row(table, [first_col_value, region, count])

            end_idx = len(table.rows) - 1

            # Merge "Severity" column cells like EC2 section
            if len(rows) > 1:
                try:
                    top_cell = table.rows[start_idx].cells[0]
                    bottom_cell = table.rows[end_idx].cells[0]
                    merged = top_cell.merge(bottom_cell)

                    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                except Exception as merge_err:
                    print(
                        f"Warning merging cells for severity '{severity_name}': {merge_err}"
                    )

        except Exception as err:
            print(f"Error processing GuardDuty severity '{severity_name}': {err}")
            continue


def add_section_cloudtrail_and_logging(doc, data):
    """
    Adds the CloudTrail & Logging configuration status section
    Columns: Parameter | Home Region | Status | Recommendation
    """

    add_title_with_number(doc, "CloudTrail & Logging", size=18)

    # Get CloudTrail check results
    global_results = data.get("global_services_scan_results", {})
    cloudtrail_data = global_results.get("cloudtrail_and_logging", {})
    findings = cloudtrail_data.get("resources_affected", [])

    # If no findings, add note and return without table
    if not findings:
        add_no_data_note(doc, "No CloudTrail or Logging configuration found")
        return

    # Create table only after confirming valid data
    headers = ["Parameter", "Home Region", "Status", "Recommendation"]
    table = create_table(doc, headers)

    # Group by parameter
    grouped = {}
    for item in findings:
        param = item.get("parameter", "-")
        region = item.get("region", "Global")
        status = item.get("status", "-")

        # Recommendation overrides
        rec = item.get("recommendation", "-")
        if param == "S3 Log Encryption":
            if status == "Enabled (KMS)":
                rec = "N/A"
            elif status == "Enabled (AES-256)":
                rec = "Upgrade to KMS encryption for enhanced key management"
        elif param in [
            "Multi-Region Logging",
            "Log File Validation",
            "CloudWatch Integration",
        ]:
            if status.lower() == "enabled":
                rec = "N/A"

        grouped.setdefault(param, []).append((region, status, rec))

    # Build table with merged Parameter column
    for param_name, rows in grouped.items():
        start_idx = len(table.rows)

        for i, (region, status, rec) in enumerate(rows):
            first_col_val = param_name if i == 0 else ""
            add_table_row(table, [first_col_val, region, status, rec])

        end_idx = len(table.rows) - 1

        # Merge parameter col if more than one region row
        if len(rows) > 1:
            try:
                top_cell = table.rows[start_idx].cells[0]
                bottom_cell = table.rows[end_idx].cells[0]
                merged = top_cell.merge(bottom_cell)
                merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as merge_err:
                print(f"Warning merging cells for '{param_name}': {merge_err}")


def add_section_iam_access_analyzer(doc, data):
    """
    Threat Detection - IAM Access Analyzer Findings
    Table: Active Findings | Region | Findings Count
    """

    add_title_with_number(doc, "Threat Detection - IAM Access Analyzer Findings", size=18)

    headers = ["Active Findings", "Region", "Findings Count"]
    table = create_table(doc, headers)

    # Group storage
    groups = {"Public Access": [], "Cross Account Access": []}

    # We need account ID for cross-account classification
    current_account_id = data.get("account_id")

    # Traverse regions
    for region_block in data.get("results", []):
        region = region_block.get("region", "-")
        region_data = region_block.get("data", {}) or {}

        analyzer_data = region_data.get("iam_access_analyzer_findings", {})
        findings = analyzer_data.get("resources_affected", [])

        # Counters
        public_count = 0
        cross_account_count = 0

        for f in findings:
            # --- PUBLIC ACCESS ---
            if f.get("isPublic", False) is True:
                # print("yes is public: ")
                public_count += 1
                continue

            # --- CROSS ACCOUNT ---
            principal = f.get("principal")
            if principal:
                try:
                    principal_str = list(principal.values())[0]
                    if "::" in principal_str:
                        principal_account = principal_str.split("::")[1].split(":")[0]
                        if principal_account != current_account_id:
                            cross_account_count += 1
                            continue
                except:
                    pass

        # Save groupings like GuardDuty section
        groups["Public Access"].append((region, str(public_count)))
        groups["Cross Account Access"].append((region, str(cross_account_count)))

    # ---- BUILD TABLE (merged first column exactly like GuardDuty) ----
    for finding_type, rows in groups.items():
        start_idx = len(table.rows)

        for i, (region, count) in enumerate(rows):
            first_col_value = finding_type if i == 0 else ""
            add_table_row(table, [first_col_value, region, count])

        end_idx = len(table.rows) - 1

        # Merge Severity / Finding Type column
        if len(rows) > 1:
            try:
                top_cell = table.rows[start_idx].cells[0]
                bottom_cell = table.rows[end_idx].cells[0]
                merged = top_cell.merge(bottom_cell)

                merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass


def add_section_observations_recommendations(doc, data):
    """
    Adds the 'Observations and Recommendations' section.
    Columns: Region | Metric | Value
    """

    add_title_with_number(doc, "Observations and Recommendations", size=18)

    headers = ["Region", "Metric", "Value"]
    table = create_table(doc, headers)

    meta_data = data.get("scanned_meta_data", [])

    if not meta_data:
        add_table_row(table, ["-", "No data available", "-"])
        return

    for region_block in meta_data:
        try:
            region = region_block.get("region", "N/A")
            region_data = region_block.get("data", {}) or {}

            total_scanned = region_data.get("total_scanned", 0)
            affected = region_data.get("affected", 0)

            # Index of where region rows start
            start_idx = len(table.rows)

            # First row (region visible)
            add_table_row(table, [region, "Total no. of Resources", str(total_scanned)])

            # Second row (region empty)
            add_table_row(table, ["", "No. of Issues Found", str(affected)])

            # Index of last row for this region
            end_idx = len(table.rows) - 1

            # Merge region column if there are 2 rows
            if end_idx > start_idx:
                try:
                    top_cell = table.rows[start_idx].cells[0]
                    bottom_cell = table.rows[end_idx].cells[0]
                    merged_cell = top_cell.merge(bottom_cell)

                    # Alignment for merged cell
                    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                except Exception as merge_err:
                    print(f"Warning: merge failed for region '{region}': {merge_err}")

        except Exception as err:
            print(f"Error processing region entry: {err}")
            continue


def add_section_recommendations_summary(doc):
    """
    Adds Recommendations Summary section with Calibri font and proper sizing.
    """

    # Title (uses your add_title() function)
    add_title_with_number(doc, "Recommendations Summary", size=18)

    # Recommendation content
    recommendations = [
        {
            "IAM:": "Enforce MFA for all users, rotate aged access keys, remove inactive identities, and regularly review unused roles and policies."
        },
        {
            "S3 Buckets:": "Restrict public access unless explicitly required, enable default encryption, and monitor access policies regularly."
        },
        {
            "EC2/RDS/Network:": "Restrict security group inbound rules to least privilege, enable encryption by default for EBS and RDS, and regularly remove unused or overly permissive rules."
        },
        {
            "Logging:": "Maintain centralized multi-region CloudTrail, enable KMS encryption for logs, forward logs to CloudWatch for real-time monitoring, and ensure log integrity validation is enabled."
        },
        {
            "General Security Best Practices:": [
                "Implement cost and usage monitoring with AWS Budgets and anomaly detection",
                "Regularly audit resource access policies for public or cross-account exposure",
                "Enable tagging standards and automate cleanup of unused or idle resources",
            ]
        },
    ]

    # Bullet creation
    def add_nested_bullets(items, level=0):
        for item in items:

            # If dictionary → key + value format
            if isinstance(item, dict):
                for key, value in item.items():

                    # Case 1: simple string recommendation
                    if isinstance(value, str):
                        para = doc.add_paragraph(
                            style=f"List Bullet{'' if level == 0 else ' ' + str(level + 1)}"
                        )

                        # Bold Key
                        run_bold = para.add_run(key + " ")
                        run_bold.bold = True
                        run_bold.font.name = "Calibri"
                        run_bold.font.size = Pt(12)

                        # Normal value
                        run_val = para.add_run(value)
                        run_val.font.name = "Calibri"
                        run_val.font.size = Pt(12)

                    # Case 2: nested list under a bold header
                    elif isinstance(value, list):
                        para = doc.add_paragraph(
                            style=f"List Bullet{'' if level == 0 else ' ' + str(level + 1)}"
                        )

                        run_bold = para.add_run(key)
                        run_bold.bold = True
                        run_bold.font.name = "Calibri"
                        run_bold.font.size = Pt(12)

                        # Process sub-bullets
                        add_nested_bullets(value, level + 1)

            # Case 3: plain string bullet
            else:
                para = doc.add_paragraph(
                    style=f"List Bullet{'' if level == 0 else ' ' + str(level + 1)}"
                )
                run = para.add_run(item)
                run.font.name = "Calibri"
                run.font.size = Pt(12)

    add_nested_bullets(recommendations)
    # add_underline_line(doc)


def add_section_conclusion(doc):
    """
    Adds conclusion section.
    """

    add_title_with_number(doc, "Conclusion", size=18)

    conclusion_text = (
        "The AWS security posture of the client's environment has been reviewed thoroughly. "
        "A number of critical and high-priority risks were identified, particularly in identity and access "
        "management, network security, and data protection. CloudThat recommends addressing the listed issues "
        "and implementing ongoing monitoring and compliance controls."
    )

    add_paragraph(doc, conclusion_text)
    # add_underline_line(doc)


# === Main Function ===


def get_report_word_function(requestData: ReportRequest):
    """
    Generate AWS Security Findings Report in Word format
    """
    global SECTION_NO
    SECTION_NO = 0

    response = get_report_from_s3_function(requestData)

    # Check response status
    if response.get("status", "") == "error":
        return response

    data = response.get("data", {})

    # === Create Document ===
    doc = Document()

    add_header_footer_images(
        doc,
        header_img_path=header_img,
        footer_img_path=footer_img,
    )

    try:
        add_report_header_section(doc=doc, data=data)
    except Exception as err:
        print(f"Error in header section: {err}")

    try:
        add_section_executive_summary(doc=doc, data=data)
    except Exception as err:
        print(f"Error in executive summary section: {err}")

    try:
        add_section_assessment_scope(doc=doc, data=data)
    except Exception as err:
        print(f"Error in assessment scope section: {err}")

    try:
        add_section_security_services_status(doc=doc, data=data)
    except Exception as err:
        print(f"Error in security services status section: {err}")

    try:
        add_section_iam(doc=doc, data=data)
    except Exception as err:
        print(f"Error in IAM section: {err}")

    try:
        add_section_s3_bucket_security(doc=doc, data=data)
    except Exception as err:
        print(f"Error in S3 bucket security section: {err}")

    try:
        add_section_ec2_network_security(doc=doc, data=data)
    except Exception as err:
        print(f"Error in EC2 network security section: {err}")

    try:
        add_section_guardduty_findings(doc=doc, data=data)
    except Exception as err:
        print(f"Error in GuardDuty Findings section: {err}")

    try:
        add_section_iam_access_analyzer(doc=doc, data=data)
    except Exception as err:
        print(f"Error in IAM Access Analyzer section: {err}")

    try:
        add_section_cloudtrail_and_logging(doc=doc, data=data)
    except Exception as err:
        print(f"Error in CloudTrail & Logging section: {err}")

    try:
        add_section_observations_recommendations(doc=doc, data=data)
    except Exception as err:
        print(f"Error in Observations and Recommendations section: {err}")

    try:
        add_section_recommendations_summary(doc=doc)
    except Exception as err:
        print(f"Error in Recommendations Summary section: {err}")

    try:
        add_section_conclusion(doc=doc)
    except Exception as err:
        print(f"Error in Conclusion section: {err}")

    # === Save to BytesIO ===
    buf = io.BytesIO()
    doc.save(buf)

    # === Response Headers ===
    account_id = requestData.account_id or ""
    account_name = data.get("account_name","")
    filename = ""

    # Case 1: Both account_id and account_name exist
    if account_id and account_name:
        filename = f"{account_id} ({account_name})_AWS_Security_Assessment_Report.docx"

    # Case 2: Only account_id exists
    elif account_id:
        filename = f"{account_id}_AWS_Security_Assessment_Report.docx"

    # Case 3: Neither exists
    else:
        filename = "AWS_Security_Assessment_Report.docx"
    

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
    }

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
