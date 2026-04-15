from fastapi.responses import Response
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from datetime import datetime, timedelta
from docx.shared import Inches, Pt, Cm, RGBColor


def add_header_footer_images(doc, header_img_path, footer_img_path):
    """
    Adds header and footer images
    """
    for section in doc.sections:
        # Set header and footer distances to zero
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0)

        # Get page and margin dimensions
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin

        #  HEADER
        header = section.header

        # Clear existing content
        for para in header.paragraphs:
            para.clear()

        # Main header image paragraph
        header_para = (
            header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        )

        # Use NEGATIVE left indent to extend beyond margin
        pf = header_para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = -left_margin  # NEGATIVE indent to start from page edge
        pf.right_indent = -right_margin  # NEGATIVE indent to extend to right edge
        pf.first_line_indent = Pt(0)

        # Add header image with full page width
        run_header = header_para.add_run()
        run_header.add_picture(header_img_path, width=page_width)

        # Partner logo paragraph
        partner_para = header.add_paragraph()
        partner_pf = partner_para.paragraph_format
        partner_pf.space_before = Pt(5)
        partner_pf.space_after = Pt(0)
        partner_pf.left_indent = Pt(0)  # Normal content margin for partner logo
        partner_pf.right_indent = Pt(0)
        partner_pf.first_line_indent = Pt(0)
        partner_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        #  FOOTER
        footer = section.footer

        # Clear existing content
        for para in footer.paragraphs:
            para.clear()

        # Footer image paragraph
        footer_para = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )

        # Use NEGATIVE indents for full width
        footer_pf = footer_para.paragraph_format
        footer_pf.space_before = Pt(0)
        footer_pf.space_after = Pt(0)
        footer_pf.left_indent = -left_margin  # NEGATIVE indent to start from page edge
        footer_pf.right_indent = (
            -right_margin
        )  # NEGATIVE indent to extend to right edge
        footer_pf.first_line_indent = Pt(0)

        # Add footer image with full page width
        run_footer = footer_para.add_run()
        run_footer.add_picture(footer_img_path, width=page_width)


# def add_title(
#     doc: Document, text: str, size: int = 24, bold: bool = True, center: bool = False, numbered: bool = True
# ):
#     """Add a title with specified font size, bold, and alignment"""
#     if numbered:
#         section_number = get_section_number(increment=True)
#         text = f"{section_number}. {text}"
#     doc.add_paragraph("")
#     para = doc.add_paragraph()
#     run = para.add_run(text)
#     run.font.name = "Calibri"
#     run.font.size = Pt(size)
#     run.bold = bold
#     if center:
#         para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# def add_title(
#     doc: Document,
#     text: str,
#     size: int = 24,
#     bold: bool = True,
#     center: bool = False,
#     numbered: bool = True,
#     level: int = 1,
# ):
#     """Add a section title (Heading 1/2/3) that supports TOC."""
#     style_name = f"Heading {level}"

#     para = doc.add_paragraph(text, style=style_name)
#     run = para.runs[0]
#     run.font.size = Pt(size)
#     run.bold = bold

#     if center:
#         para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_title(
    doc: Document,
    text: str,
    font_name: str = "Calibri",
    font_color: str = "000000",
    size: int = 24,
    bold: bool = True,
    center: bool = False,
    numbered: bool = True,
    level: int = 1,
):
    """Add a section title (Heading 1/2/3) that supports TOC and indentation."""

    style_name = f"Heading {level}"

    para = doc.add_paragraph(text, style=style_name)
    run = para.runs[0]
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(font_color)

    # ---- Align if needed ----
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Indentation based on heading level ----
    # pf = para.paragraph_format

    # if level == 1:
    #     pf.left_indent = Inches(0)
    # elif level == 2:
    #     pf.left_indent = Inches(0.3)
    # elif level == 3:
    #     pf.left_indent = Inches(0.6)
    # elif level == 4:
    #     pf.left_indent = Inches(0.9)
    # else:
    #     pf.left_indent = Inches(0)

    return para


def add_key_value(doc: Document, key: str, value: str):
    """Add key-value paragraph where key is bold, value normal (Calibri 12)"""
    para = doc.add_paragraph()
    run_key = para.add_run(f"{key}: ")
    run_key.font.name = "Calibri"
    run_key.font.size = Pt(12)
    run_key.bold = True

    run_value = para.add_run(value)
    run_value.font.name = "Calibri"
    run_value.font.size = Pt(12)


def add_bullet_points(doc: Document, items: list):
    """Add bullet points in Calibri 12"""
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(12)


def add_table_row(table, values: list):
    """
    Add a row to a table with dynamic columns.
    """
    row = table.add_row().cells

    for i, val in enumerate(values):
        if i >= len(row):
            break  # Safety if more values than columns
        row[i].text = val

        # Set font
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(12)
        row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in row[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply cell borders
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
        tcPr.append(tcBorders)


def add_paragraph(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)


def create_table(doc, headers):
    """
    Create a table with visible borders and centered alignment.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set header text and formatting
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Make borders visible
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._tc.append(tcPr)
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Center all header cell contents horizontally and vertically
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table


def add_bullet(doc, text):
    """
    Adds a bulleted paragraph to the document.
    """
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(11)


# def add_underline_line(doc: Document, length: int = 40):
#     doc.add_paragraph("_" * length)


def add_underline_line(
    doc, line_width=0.5, space_before=6, space_after=6, line_color="BBBBBB"
):
    """
    Adds a thin, light horizontal line
    """
    p = doc.add_paragraph()

    # Set spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # Create paragraph border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")  # single line
    bottom.set(qn("w:sz"), str(int(line_width * 8)))  # thickness in eighths of a point
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), line_color)  # lighter color
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_centered_text(doc, text, size=22, bold=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
