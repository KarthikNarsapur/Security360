from fastapi.responses import Response
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
import io
from datetime import datetime, timedelta
from docx.shared import Inches, Pt, Cm
from Model.model import ReportRequest
from utils.upload_to_s3 import get_report_from_s3_function
from word_report_generator.utils import (
    add_header_footer_images,
    add_title,
    add_key_value,
    add_bullet_points,
    add_table_row,
    add_paragraph,
    create_table,
    add_bullet,
    add_underline_line,
    add_centered_text,
)

ASSETS_DIR = "assets/images"
header_img = f"{ASSETS_DIR}/full_header.png"
footer_img = f"{ASSETS_DIR}/footer.png"


def extract_pillar_from_id(check_id):
    """
    Convert check ID prefix into pillar name.
    Example: COST05-BP01 → Cost Optimization
    """

    if not isinstance(check_id, str):
        return ""

    prefix = check_id.split("-")[0].lower()

    if prefix.startswith("sec"):
        return "Security"
    if prefix.startswith("cost"):
        return "Cost Optimization"
    if prefix.startswith("ops"):
        return "Operational Excellence"
    if prefix.startswith("rel"):
        return "Reliability"
    if prefix.startswith("perf"):
        return "Performance Efficiency"
    if prefix.startswith("sus"):
        return "Sustainability"

    return ""


def process_awarf_data(data):
    """
    Process WAR review data and extract useful summary information.
    """

    results = data.get("results", [])
    pillars = data.get("pillars", [])

    processed = {
        "pillars": {},
        "totals": {
            "passed": 0,
            "failed": 0,
            "not_available": 0,
            "total_scanned": 0,
            "total_affected": 0,
        },
        "high_risk_items": [],
        "medium_risk_items": [],
        "low_risk_items": [],
        "failed_checks": [],
        "passed_checks": [],
        "not_available_checks": [],
    }

    # ---- Initialize per-pillar bucket ----
    for pillar in pillars:
        processed["pillars"][pillar] = {
            "total_checks": 0,
            "passed": 0,
            "failed": 0,
            "not_available": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
            "total_scanned": 0,
            "total_affected": 0,
        }

    # ---- Iterate through each region result block ----
    for region_block in results:
        region_data = region_block.get("data", {})

        for key, check in region_data.items():
            if check is None:
                print("yes check is none ",key)
                continue
            severity = check.get("severity_level", "")
            status = check.get("status", "")
            pillar_name = extract_pillar_from_id(check.get("id", ""))

            # Skip checks not belonging to declared pillars
            if pillar_name not in processed["pillars"]:
                continue

            pillar_summary = processed["pillars"][pillar_name]
            pillar_summary["total_checks"] += 1

            total_scanned = check.get("additional_info", {}).get("total_scanned", 0)
            affected = check.get("additional_info", {}).get("affected", 0)

            pillar_summary["total_scanned"] += total_scanned
            pillar_summary["total_affected"] += affected

            # ---- Count statuses ----
            if status == "passed":
                pillar_summary["passed"] += 1
                processed["totals"]["passed"] += 1
                processed["passed_checks"].append(check)

            elif status == "failed":
                pillar_summary["failed"] += 1
                processed["totals"]["failed"] += 1
                processed["failed_checks"].append(check)

            elif status == "not_available":
                pillar_summary["not_available"] += 1
                processed["totals"]["not_available"] += 1
                processed["not_available_checks"].append(check)

            # ---- Count severity ----
            sev = severity.lower()
            if sev == "high":
                pillar_summary["high"] += 1
                processed["high_risk_items"].append(check)
            elif sev == "medium":
                pillar_summary["medium"] += 1
                processed["medium_risk_items"].append(check)
            elif sev == "low":
                pillar_summary["low"] += 1
                processed["low_risk_items"].append(check)

            # ---- Global totals ----
            processed["totals"]["total_scanned"] += total_scanned
            processed["totals"]["total_affected"] += affected

    return processed


def add_cover_page(doc, data):
    section = doc.sections[0]

    # ---- Remove header/footer from the cover page ----
    section.different_first_page_header_footer = True
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False

    # ---- Add AWS Well-Architected Icon ----
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run()
    run.add_picture("assets/images/awarf.png", width=Inches(1.2))

    # ---- Add Title (matching your provided image) ----
    title_lines = [
        "AWS Well-Architected",
        "Tool WAR REVIEW",
        "ARCHITECTURE 2025 -",
        "AWS Well-Architected",
        "Framework Report",
    ]

    for line in title_lines:
        add_centered_text(doc, line, size=28, bold=True)

    # ---- Add AWS Account ID ----
    doc.add_paragraph("")  # spacing
    pid = doc.add_paragraph()
    pid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pid.add_run(f"AWS Account ID: {data.get("account_id","")}")
    run.font.size = Pt(16)
    run.font.name = "Calibri"

    # ---- Page Break ----
    # doc.add_page_break()


def add_table_of_contents(doc):
    """
    Inserts a dynamic table of contents that Word will update automatically.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fldSimple)

    doc.add_page_break()


def add_workload_properties_section(doc, data):
    """
    Generates the 'Workload properties' section exactly like screenshot.
    """

    # ---- Section Title (Heading 1 for TOC) ----
    add_title(
        doc,
        "Workload properties",
        size=24,
        bold=True,
        center=False,
        numbered=False,
        level=1,
    )

    # ---- Workload name ----
    add_key_value(
        doc, "Workload name", data.get("workload_name", "WAR REVIEW ARCHITECTURE")
    )

    # ---- Description ----
    pillar_count = len(data.get("pillars", []))
    description_text = f"{pillar_count} Pillar Review"
    add_key_value(doc, "Description", description_text)

    # ---- Review owner (empty as you required) ----
    add_key_value(doc, "Review owner", "")

    # ---- Industry type ----
    # add_key_value(doc, "Industry type", data.get("industry_type", "-"))

    # ---- Industry ----
    # add_key_value(doc, "Industry", data.get("industry", "-"))

    # ---- Environment ----
    # add_key_value(doc, "Environment", data.get("environment", "-"))

    # ---- AWS Regions ----
    regions = data.get("regions", [])
    regions_string = ", ".join(regions) if regions else "-"
    add_key_value(doc, "AWS Regions", regions_string)

    # ---- Non-AWS regions ----
    # add_key_value(doc, "Non-AWS regions", "-")

    # ---- Account IDs ----
    acc_id = data.get("account_id", "")
    add_key_value(doc, "Account IDs", acc_id)

    # ---- Architectural design ----
    # add_key_value(doc, "Architectural design", "")


def add_lens_overview_section(doc, data, processed):
    # ---- Section Title ----
    add_title(doc, "Lens overview", size=24, bold=True, numbered=False, level=1)

    # ---- Questions answered summary ----
    total_passed = processed["totals"]["passed"]
    total_failed = processed["totals"]["failed"]
    total_na = processed["totals"]["not_available"]

    total_answered = total_passed + total_failed
    total_questions = total_passed + total_failed + total_na

    add_key_value(doc, "Questions answered", f"{total_answered}/{total_questions}")

    # ---- Pillar question breakdown table ----
    table = create_table(doc, ["Pillar", "Questions answered"])

    for pillar, summary in processed["pillars"].items():
        p_passed = summary["passed"]
        p_failed = summary["failed"]
        p_na = summary["not_available"]

        p_answered = p_passed + p_failed
        p_total = p_passed + p_failed + p_na

        add_table_row(table, [pillar, f"{p_answered}/{p_total}"])

    # ---- Lens notes ----
    # add_key_value(doc, "Lens notes", "-")


def create_rounded_box(doc):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Padding
    cell.margin_top = Cm(0.2)
    cell.margin_bottom = Cm(0.2)
    cell.margin_left = Cm(0.3)
    cell.margin_right = Cm(0.3)

    # Border style
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")

    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")  # border thickness
        # border.set(qn("w:color"), "4A90E2")  # light blue
        borders.append(border)

    tcPr.append(borders)

    return cell

def add_hyperlink(paragraph, url, text):
    # Create relationship in document
    part = paragraph.part
    r_id = part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )

    # Create the hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create formatting run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Styling (blue + underline)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    # Text element
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)

    # Append to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


# def add_each_check_summary_section(doc, data):
#     add_title(doc, "Check summary", size=24, bold=True, numbered=False, level=1)

#     sr = 1

#     status_map = {
#         "not_available": "Not Available",
#         "failed": "Non Compliant",
#         "passed": "Compliant",
#     }

#     for region_block in data.get("results", []):
#         region_data = region_block.get("data", {})

#         for key, check in region_data.items():
            
#             if check is None:
#                 continue

#             cell = create_rounded_box(doc)  # rounded box container
#             box_para = cell.add_paragraph()
            
#                  # ---- Check ID + check title ----
#             check_id = check.get("id", "")
#             check_name = check.get("check_name", "")

#             run = box_para.add_run(f"{check_id}. {check_name}")
#             run.bold = True
#             run.font.size = Pt(14)

#             # Status
#             raw_status = check.get("status", "").lower()
#             mapped_status = status_map.get(raw_status, "Unknown")
#             p1 = cell.add_paragraph()
#             r1 = p1.add_run("Status: ")
#             r1.bold = True
#             p1.add_run(mapped_status)

#             # Severity
#             p2 = cell.add_paragraph()
#             r2 = p2.add_run("Severity: ")
#             r2.bold = True
#             p2.add_run(check.get("severity_level", "-"))

#             # Additional Info
#             add_info = check.get("additional_info", {})
#             total_scanned = add_info.get("total_scanned", 0)
#             affected = add_info.get("affected", 0)

#             p3 = cell.add_paragraph()
#             r3 = p3.add_run("Total scanned: ")
#             r3.bold = True
#             p3.add_run(str(total_scanned))

#             p4 = cell.add_paragraph()
#             r4 = p4.add_run("Affected: ")
#             r4.bold = True
#             p4.add_run(str(affected))

#             # ---- AWS Doc Link ----
#             # doc_link = check.get("aws_doc_link", "")
#             # if doc_link:
#             #     p5 = cell.add_paragraph()
#             #     r5 = p5.add_run("AWS Doc: ")
#             #     r5.bold = True
#             #     link_run = p5.add_run(doc_link)
#             #     link_run.font.color.rgb = RGBColor(0, 0, 255)
#             #     link_run.font.underline = True

#             doc_link = check.get("aws_doc_link", "")
#             if doc_link:
#                 p5 = cell.add_paragraph()
#                 r5 = p5.add_run("AWS Doc: ")
#                 r5.bold = True

#                 # professional clickable text
#                 link_text = "AWS Documentation"

#                 hyperlink = add_hyperlink(p5, doc_link, link_text)

#             doc.add_paragraph("")  # spacing

#             sr += 1





def add_each_check_summary_section(doc, data):
    add_title(doc, "Check summary", size=24, bold=True, numbered=False, level=1)

    status_map = {
        "not_available": "Not Available",
        "failed": "Non Compliant",
        "passed": "Compliant",
    }

    # Sorting helper
    def sort_check_id(check_id):
        # Example: OPS08-BP03
        try:
            prefix, bp = check_id.split("-")
            pillar_num = int(prefix[-2:])
            bp_num = int(bp.replace("BP", ""))
            pillar_str = prefix[:-2]
            return (pillar_str, pillar_num, bp_num)
        except:
            return ("ZZZ", 999, 999)  # fallback: push unknowns to bottom

    for region_block in data.get("results", []):
        
        region = region_block.get("region", "Unknown Region")
        region_data = region_block.get("data", {})

        add_title(doc, f"Region: {region}", size=16, bold=True, numbered=False, level=2)

        # Group by pillar
        pillar_map = {}

        for key, check in region_data.items():
            if check is None:
                continue

            pillar = extract_pillar_from_id(check.get("id", "")) or "Unknown Pillar"

            if pillar not in pillar_map:
                pillar_map[pillar] = []

            pillar_map[pillar].append(check)

        # Iterate pillar groups
        for pillar, checks in pillar_map.items():

            add_title(doc, f"Pillar: {pillar}", size=14, bold=True, numbered=False, level=3)

            # ---- SORT CHECKS BY ID ----
            checks_sorted = sorted(checks, key=lambda c: sort_check_id(c.get("id", "")))

            for check in checks_sorted:

                cell = create_rounded_box(doc)
                box_para = cell.add_paragraph()

                # Check ID + title
                run = box_para.add_run(f"{check.get('id','')}. {check.get('check_name','')}")
                run.bold = True
                run.font.size = Pt(14)

                # Status
                raw_status = check.get("status", "").lower()
                mapped_status = status_map.get(raw_status, "Unknown")
                p1 = cell.add_paragraph()
                r1 = p1.add_run("Status: ")
                r1.bold = True
                p1.add_run(mapped_status)

                # Severity
                p2 = cell.add_paragraph()
                r2 = p2.add_run("Severity: ")
                r2.bold = True
                p2.add_run(check.get("severity_level", "-"))

                # Additional Info
                add_info = check.get("additional_info", {})
                total_scanned = add_info.get("total_scanned", 0)
                affected = add_info.get("affected", 0)

                p3 = cell.add_paragraph()
                r3 = p3.add_run("Total scanned: ")
                r3.bold = True
                p3.add_run(str(total_scanned))

                p4 = cell.add_paragraph()
                r4 = p4.add_run("Affected: ")
                r4.bold = True
                p4.add_run(str(affected))

                # AWS Documentation Link
                doc_link = check.get("aws_doc_link", "")
                if doc_link:
                    p5 = cell.add_paragraph()
                    r5 = p5.add_run("AWS Doc: ")
                    r5.bold = True
                    add_hyperlink(p5, doc_link, "AWS Documentation")

                doc.add_paragraph("")



def get_awarf_report_word_function(requestData: ReportRequest):
    """
    Generate AWS Security Findings Report in Word format
    """
    global SECTION_NO
    SECTION_NO = 0

    response = get_report_from_s3_function(requestData)

    # Check response status
    if response.get("status", "") == "error":
        return response

    data = response.get("data", {})

    # === Create Document ===
    doc = Document()

    add_header_footer_images(
        doc,
        header_img_path=header_img,
        footer_img_path=footer_img,
    )

    try:
        processed_data = process_awarf_data(data)
    except Exception as err:
        print(f"Error in processing data: {err}")

    try:
        add_cover_page(doc=doc, data=data)
    except Exception as err:
        print(f"Error in header section: {err}")

    try:
        add_table_of_contents(doc=doc)
    except Exception as err:
        print(f"Error in table of content section: {err}")

    try:
        add_workload_properties_section(doc=doc, data=data)
    except Exception as err:
        print(f"Error in workload properties: {err}")

    try:
        add_lens_overview_section(doc=doc, data=data, processed=processed_data)
    except Exception as err:
        print(f"Error in lens overview : {err}")

    try:
        add_each_check_summary_section(doc=doc, data=data)
    except Exception as err:
        print(f"Error in each check summary : {err}")

    # === Save to BytesIO ===
    buf = io.BytesIO()
    doc.save(buf)

    # === Response Headers ===
    now_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    filename = ""
    if requestData.account_id:
        filename = f"{requestData.account_id}_{now_str}.docx"
    else:
        filename = f"{now_str}.docx"

    # if requestData.account_id:
    #     filename = f"{requestData.account_id}_AWS_Well-Architected_Framework_Report_{now_str}.docx"
    # else:
    #     filename = f"AWS_Well-Architected_Framework_Report_{now_str}.docx"

    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
    }

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
